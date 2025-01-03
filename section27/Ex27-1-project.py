import tkinter as tk
from tkinter import scrolledtext, ttk, filedialog
from docx import Document
from datetime import datetime
from openai import OpenAI
import threading
import os
import base64
from PIL import Image, ImageTk
import requests
from io import BytesIO
import sv_ttk  # 모던한 UI를 위한 Sun Valley ttk 테마
import time

class MultiAIApp:
    def __init__(self, root):
        # 초기화 메서드: 앱의 UI 및 기능 초기화
        self.root = root
        self.root.title("AI Assistant")
        self.root.geometry("1000x800")

        # 테마 설정
        sv_ttk.set_theme("dark")  # 다크 테마 적용

        # 스타일 설정
        self.style = ttk.Style()
        self.style.configure('Chat.TFrame', padding=10)
        self.style.configure('Control.TFrame', padding=5)
        self.style.configure('Title.TLabel', font=('Helvetica', 12, 'bold'))
        self.style.configure('Status.TLabel', font=('Helvetica', 10))

        # OpenAI 클라이언트 및 대화 기록 초기화
        self.client = OpenAI()
        self.conversation_history = []

        # UI 요소 초기화
        self.current_photo = None  # 이미지 생성 결과 참조 유지용
        self.chat_area = None  # 채팅 영역
        self.message_entry = None  # 메시지 입력 영역
        self.loading_label = None  # 상태 표시 라벨
        self.image_prompt = None  # 이미지 생성 프롬프트 입력 필드
        self.image_label = None  # 생성된 이미지 표시 라벨
        self.image_status = None  # 이미지 상태 메시지 라벨
        self.tts_text = None  # TTS 입력 필드
        self.stt_result = None  # STT 변환 결과 출력 필드

        # UI 설정 메서드 호출
        self.setup_ui()

    def setup_ui(self):
        # UI 초기화 및 구성
        main_container = ttk.Frame(self.root, padding="10")
        main_container.pack(fill='both', expand=True)

        # 노트북(탭) 생성
        self.notebook = ttk.Notebook(main_container)
        self.notebook.pack(fill='both', expand=True)

        # 탭 설정
        self.setup_chat_tab()  # 채팅 탭 구성
        self.setup_image_tab()  # 이미지 생성 탭 구성
        self.setup_voice_tab()  # 음성 변환 탭 구성


    def setup_chat_tab(self):
        # 채팅 탭 UI 구성
        self.chat_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.chat_frame, text=' 💬 채팅 ')

        # 타이틀 프레임
        title_frame = ttk.Frame(self.chat_frame)
        title_frame.pack(fill='x', pady=(0, 10))
        ttk.Label(title_frame, text="AI와의 대화", style='Title.TLabel').pack(side='left')

        # 채팅 영역 (테두리와 패딩 추가)
        chat_container = ttk.Frame(self.chat_frame, relief='solid', borderwidth=1)
        chat_container.pack(fill='both', expand=True, pady=(0, 10))

        self.chat_area = scrolledtext.ScrolledText(
            chat_container,             # 채팅 컨테이너에 위젯 배치
            wrap=tk.WORD,               # 단어 단위로 자동 줄바꿈
            height=20,                  # 채팅창 높이 설정 (행 단위)
            font=('Helvetica', 10),     # 폰트 설정: Helvetica, 크기 10
            bg='#000000'                # 배경색 설정: 검정색
        )
        self.chat_area.pack(fill='both', expand=True, padx=5, pady=5)

        # 입력 영역 컨테이너
        input_container = ttk.Frame(self.chat_frame, style='Control.TFrame')
        input_container.pack(fill='x', pady=(0, 5))

        # 메시지 입력 프레임
        input_frame = ttk.Frame(input_container)
        input_frame.pack(fill='x', pady=5)

        self.message_entry = ttk.Entry(
            input_frame,  # 입력창을 input_frame에 배치
            font=('Helvetica', 10)  # 폰트 설정: Helvetica, 크기 10
        )
        self.message_entry.pack(
            side='left',  # 왼쪽 정렬
            fill='x',  # x축 방향으로 공간 채우기
            expand=True,  # 남는 공간을 채우도록 확장
            padx=(0, 5)  # 좌우 패딩: 왼쪽 0, 오른쪽 5픽셀
        )
        self.message_entry.bind(
            '<Return>',  # Enter 키 이벤트 바인딩
            self.send_message  # Enter 키 누르면 send_message 함수 실행
        )

        ttk.Button(
            input_frame,
            text="전송",
            style='Accent.TButton',
            command=self.send_message
        ).pack(side='right')

        # 컨트롤 버튼 프레임
        control_frame = ttk.Frame(input_container)
        control_frame.pack(fill='x', pady=5)

        ttk.Button(
            control_frame,
            text="💾 대화 내역 저장",
            command=self.save_to_docx
        ).pack(side='left', padx=2)

        ttk.Button(
            control_frame,
            text="🔄 새 채팅",
            command=self.new_chat
        ).pack(side='left', padx=2)

        ttk.Button(
            control_frame,
            text="🔊 음성으로 듣기",
            command=self.speak_last_response
        ).pack(side='left', padx=2)

        # 상태 표시 레이블
        self.loading_label = ttk.Label(
            self.chat_frame,
            text="",
            style='Status.TLabel'
        )
        self.loading_label.pack(pady=5)

    def send_message(self, event=None):
        # 메시지 전송 메서드
        message = self.message_entry.get().strip()  # 입력 필드에서 메시지 가져오기
        if not message:
            return

        # 입력 필드 초기화 및 사용자 메시지 표시
        self.message_entry.delete(0, tk.END)
        self.chat_area.insert(tk.END, f"나: {message}\n")
        self.chat_area.see(tk.END)

        # 응답 생성 중 표시
        self.loading_label.config(text="🤔 응답을 생성하는 중...")
        threading.Thread(target=self.get_gpt_response, args=(message,)).start()  # 응답 생성 비동기 처리

    def get_gpt_response(self, message):
        # OpenAI GPT 모델에 사용자 메시지 전달 및 응답 받기
        try:
            self.conversation_history.append({"role": "user", "content": message})

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=self.conversation_history
            )

            assistant_message = response.choices[0].message.content  # GPT의 응답 가져오기
            self.conversation_history.append(
                {"role": "assistant", "content": assistant_message}
            )

            # UI 업데이트 요청
            self.root.after(0, self.update_chat_area, assistant_message)
        except Exception as e:
            self.root.after(0, self.update_chat_area, f"❌ Error: {str(e)}")
        finally:
            self.root.after(0, self.loading_label.config, {"text": ""})

    def update_chat_area(self, response_text):
        # GPT 응답을 채팅 영역에 표시
        self.chat_area.insert(tk.END, f"GPT: {response_text}\n\n")
        self.chat_area.see(tk.END)

    def new_chat(self):
        # 새 대화 시작 메서드
        if self.chat_area.get("1.0", tk.END).strip():
            self.auto_save_chat()  # 이전 대화 자동 저장

        self.conversation_history = []  # 대화 기록 초기화
        self.chat_area.delete("1.0", tk.END)
        self.chat_area.insert(tk.END, "✨ 새로운 대화가 시작되었습니다.\n\n")
        self.chat_area.see(tk.END)

    def auto_save_chat(self):
        # 대화 자동 저장 메서드
        save_dir = "chat_history"
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)

        filename = f"대화내역_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(save_dir, filename)

        self.save_chat_to_file(file_path)
        self.chat_area.insert(tk.END, f"💾 대화 내역이 자동 저장되었습니다: {file_path}\n\n")
        self.chat_area.see(tk.END)

    def save_chat_to_file(self, file_path):
        # 대화 내용을 Word 파일로 저장
        doc = Document()
        doc.add_heading('AI Assistant 대화 내역', 0)
        doc.add_paragraph(
            f"저장 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}"
        )
        doc.add_paragraph('=' * 50)

        chat_content = self.chat_area.get("1.0", tk.END)
        paragraphs = chat_content.split('\n')

        for para in paragraphs:
            if para.strip():
                if para.startswith('나:'):
                    p = doc.add_paragraph()
                    p.add_run('나: ').bold = True
                    p.add_run(para[3:])
                elif para.startswith('GPT:'):
                    p = doc.add_paragraph()
                    p.add_run('GPT: ').bold = True
                    p.add_run(para[4:])
                else:
                    doc.add_paragraph(para)

        doc.save(file_path)

    def save_to_docx(self):
        # 대화 내역을 Word 파일로 저장하기 위한 메서드
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=f"대화내역_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if not file_path:
                return

            self.save_chat_to_file(file_path)
            self.chat_area.insert(tk.END, f"💾 대화 내역이 저장되었습니다: {file_path}\n\n")
            self.chat_area.see(tk.END)

        except Exception as e:
            self.chat_area.insert(tk.END,
                                  f"❌ Error: 대화 내역 저장 중 오류가 발생했습니다. {str(e)}\n\n")
            self.chat_area.see(tk.END)


    def setup_image_tab(self):
        # 이미지 생성 탭 UI 구성
        self.image_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.image_frame, text=' 🎨 이미지 생성 ')

        # 타이틀
        ttk.Label(
            self.image_frame,
            text="AI 이미지 생성",
            style='Title.TLabel'
        ).pack(pady=(0, 10))

        # 컨트롤 패널
        control_panel = ttk.LabelFrame(self.image_frame, text="이미지 설정", padding=10)
        control_panel.pack(fill='x', padx=5, pady=5)

        # 프롬프트 입력
        prompt_frame = ttk.Frame(control_panel)
        prompt_frame.pack(fill='x', pady=5)

        ttk.Label(prompt_frame, text="이미지 설명:").pack(side='left', padx=(0, 5))
        self.image_prompt = ttk.Entry(prompt_frame, width=50)
        self.image_prompt.pack(side='left', fill='x', expand=True)

        # 이미지 설정
        settings_frame = ttk.Frame(control_panel)
        settings_frame.pack(fill='x', pady=10)

        # 크기 선택
        size_frame = ttk.LabelFrame(settings_frame, text="크기", padding=5)
        size_frame.pack(side='left', padx=5)

        self.size_var = tk.StringVar(value="1024x1024")
        ttk.Radiobutton(
            size_frame,
            text="정사각형",
            variable=self.size_var,
            value="1024x1024"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            size_frame,
            text="세로형",
            variable=self.size_var,
            value="1024x1792"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            size_frame,
            text="가로형",
            variable=self.size_var,
            value="1792x1024"
        ).pack(side='left', padx=5)

        # 품질 선택
        quality_frame = ttk.LabelFrame(settings_frame, text="품질", padding=5)
        quality_frame.pack(side='left', padx=5)

        self.quality_var = tk.StringVar(value="standard")
        ttk.Radiobutton(
            quality_frame,
            text="일반",
            variable=self.quality_var,
            value="standard"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            quality_frame,
            text="고품질",
            variable=self.quality_var,
            value="hd"
        ).pack(side='left', padx=5)

        # 생성 버튼
        ttk.Button(
            control_panel,
            text="🎨 이미지 생성",
            style='Accent.TButton',
            command=self.generate_image
        ).pack(pady=10)

        # 이미지 표시 영역
        image_display = ttk.LabelFrame(self.image_frame, text="생성된 이미지", padding=10)
        image_display.pack(fill='both', expand=True, padx=5, pady=5)

        self.image_label = ttk.Label(image_display)
        self.image_label.pack(pady=10)

        self.image_status = ttk.Label(
            image_display,
            text="",
            style='Status.TLabel'
        )
        self.image_status.pack(pady=5)

    def setup_voice_tab(self):
        # 음성 변환 탭 UI 구성
        self.voice_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.voice_frame, text=' 🎤 음성 변환 ')

        # TTS 섹션
        tts_frame = ttk.LabelFrame(
            self.voice_frame,
            text="텍스트 → 음성 변환",
            padding=10
        )
        tts_frame.pack(fill='x', padx=5, pady=5)

        ttk.Label(tts_frame, text="변환할 텍스트:").pack(pady=2)

        self.tts_text = scrolledtext.ScrolledText(
            tts_frame,
            height=4,
            font=('Helvetica', 10)
        )
        self.tts_text.pack(fill='x', pady=5)

        # 음성 설정
        voice_frame = ttk.Frame(tts_frame)
        voice_frame.pack(fill='x', pady=5)

        ttk.Label(voice_frame, text="음성 선택:").pack(side='left', padx=5)
        self.voice_var = tk.StringVar(value="alloy")
        voices = ["alloy", "echo", "fable", "onyx", "nova", "shimmer"]
        voice_menu = ttk.OptionMenu(voice_frame, self.voice_var, "alloy", *voices)
        voice_menu.pack(side='left', padx=5)

        ttk.Button(
            tts_frame,
            text="🔊 음성 생성",
            style='Accent.TButton',
            command=self.generate_speech
        ).pack(pady=5)

        # STT 섹션
        stt_frame = ttk.LabelFrame(
            self.voice_frame,
            text="음성 → 텍스트 변환",
            padding=10
        )
        stt_frame.pack(fill='x', padx=5, pady=10)

        # 컨트롤
        control_frame = ttk.Frame(stt_frame)
        control_frame.pack(fill='x', pady=5)

        ttk.Button(
            control_frame,
            text="🎤 음성 파일 선택",
            command=self.select_audio_file
        ).pack(side='left', padx=5)

        # 언어 선택
        self.lang_var = tk.StringVar(value="auto")
        langs = ["auto", "ko", "en", "ja"]
        ttk.Label(control_frame, text="언어:").pack(side='left', padx=5)
        lang_menu = ttk.OptionMenu(control_frame, self.lang_var, "auto", *langs)
        lang_menu.pack(side='left', padx=5)

        # 결과 표시
        self.stt_result = scrolledtext.ScrolledText(
            stt_frame,
            height=4,
            font=('Helvetica', 10)
        )
        self.stt_result.pack(fill='x', pady=5)

    def generate_image(self):
        # 이미지 생성 메서드
        prompt = self.image_prompt.get().strip()
        if not prompt:
            self.image_status.config(text="⚠️ 이미지 설명을 입력해주세요.")
            return

        self.image_status.config(text="🎨 이미지를 생성하는 중...")
        threading.Thread(target=self.generate_image_thread, args=(prompt,)).start()

    def generate_image_thread(self, prompt):
        try:
            # DALL-E API를 호출하여 이미지 생성 요청
            response = self.client.images.generate(
                model="dall-e-3",  # 사용할 DALL-E 모델
                prompt=prompt,  # 사용자 입력 프롬프트
                size=self.size_var.get(),  # 이미지 크기 (UI에서 선택된 값)
                quality=self.quality_var.get(),  # 이미지 품질 (UI에서 선택된 값)
                n=1,  # 생성할 이미지 개수
                style="natural"  # 이미지 스타일
            )

            # 생성된 이미지의 URL 가져오기
            image_url = response.data[0].url
            image_response = requests.get(image_url)
            image = Image.open(BytesIO(image_response.content))  # URL에서 이미지 읽기

            # 이미지 크기 조정 (최대 800x800 픽셀)
            max_size = (800, 800)
            image.thumbnail(max_size, Image.Resampling.LANCZOS)

            photo = ImageTk.PhotoImage(image)  # Tkinter용 이미지 객체 생성

            # UI에 이미지와 상태 메시지 업데이트
            self.root.after(0, self.update_image, photo, "✨ 이미지가 생성되었습니다!")

            # 생성된 이미지를 로컬 폴더에 자동 저장
            save_dir = "generated_images"  # 저장 디렉토리
            if not os.path.exists(save_dir):
                os.makedirs(save_dir)  # 디렉토리가 없으면 생성

            # 파일 이름에 타임스탬프 추가하여 저장
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            save_path = os.path.join(save_dir, f"generated_image_{timestamp}.png")
            image.save(save_path)

        except Exception as e:
            # 예외 발생 시 오류 메시지 출력 및 UI 업데이트
            error_message = str(e)
            self.root.after(0, self.image_status.config,
                            {"text": f"❌ Error: {error_message}"})
            print(f"Error in image generation: {error_message}")

    def update_image(self, photo, status_text):
        # UI에서 이미지를 업데이트하고 상태 텍스트를 변경
        self.current_photo = photo  # Tkinter에서 이미지가 삭제되지 않도록 참조 유지
        self.image_label.config(image=photo)  # 이미지 레이블 업데이트
        self.image_status.config(text=status_text)  # 상태 메시지 업데이트

    def generate_speech(self):
        # 텍스트를 받아서 음성으로 변환
        text = self.tts_text.get("1.0", tk.END).strip()  # 텍스트 상자의 내용을 가져옴
        if not text:  # 텍스트가 없으면 종료
            return

        try:
            # 텍스트를 음성으로 변환하는 API 호출
            response = self.client.audio.speech.create(
                model="tts-1",  # TTS 모델
                voice=self.voice_var.get(),  # 선택된 음성
                input=text  # 입력 텍스트
            )

            # 음성 파일을 임시 파일로 저장
            temp_file = "temp_speech.mp3"
            response.stream_to_file(temp_file)

            # 파일 저장 위치를 사용자에게 선택받기
            save_path = filedialog.asksaveasfilename(
                defaultextension=".mp3",  # 기본 확장자
                filetypes=[("MP3 files", "*.mp3")],  # 파일 형식
                initialfile=f"speech_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3"
            )

            if save_path:
                # 임시 파일을 선택한 위치로 이동
                if os.path.exists(temp_file):
                    os.replace(temp_file, save_path)

        except Exception as e:
            # 예외 발생 시 임시 파일 삭제 및 오류 처리
            if os.path.exists(temp_file):
                os.remove(temp_file)
            raise e

    def speak_last_response(self):
        # 채팅에서 마지막 GPT 응답을 음성으로 변환
        chat_content = self.chat_area.get("1.0", tk.END)  # 채팅 내용 가져오기
        lines = chat_content.split('\n')
        last_response = None

        # 마지막 GPT 응답 텍스트 추출
        for line in reversed(lines):
            if line.startswith("GPT: "):  # GPT 응답인 줄 확인
                last_response = line[5:]  # "GPT: " 제거
                break

        if last_response:
            # 텍스트를 음성 변환 입력 상자에 삽입
            self.tts_text.delete("1.0", tk.END)
            self.tts_text.insert("1.0", last_response)
            self.generate_speech()  # 음성 생성 호출

    def select_audio_file(self):
        # 사용자가 음성 파일을 선택하도록 파일 탐색기 열기
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Audio files", "*.mp3 *.mp4 *.mpeg *.mpga *.m4a *.wav *.webm")  # 지원 파일 형식
            ]
        )

        if not file_path:  # 파일을 선택하지 않은 경우 종료
            return

        try:
            # 선택된 파일을 열어 STT API 호출
            with open(file_path, "rb") as audio_file:
                language = self.lang_var.get()  # 언어 선택 값 가져오기
                if language == "auto":  # 자동 언어 감지 설정
                    language = None

                # 음성 파일을 텍스트로 변환
                transcript = self.client.audio.transcriptions.create(
                    model="whisper-1",  # 음성 인식 모델
                    file=audio_file,  # 음성 파일
                    language=language  # 선택한 언어
                )

                # 변환 결과를 UI의 텍스트 상자에 표시
                self.stt_result.delete("1.0", tk.END)
                self.stt_result.insert("1.0", transcript.text)

        except Exception as e:
            # 예외 발생 시 오류 메시지 표시
            self.stt_result.delete("1.0", tk.END)
            self.stt_result.insert("1.0", f"❌ Error: {str(e)}")


if __name__ == "__main__":
    # 프로그램 실행 진입점
    root = tk.Tk()  # Tkinter의 메인 윈도우 생성
    app = MultiAIApp(root)  # MultiAIApp 클래스의 인스턴스 생성
    root.mainloop()  # Tkinter 이벤트 루프 시작